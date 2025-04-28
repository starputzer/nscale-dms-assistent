
import os
import pandas as pd
from datetime import datetime
import magic
import fitz  # PyMuPDF für PDF-Metadaten
import docx  # für Word-Dokumente
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional, Union

class DocumentScanner:
    """Scanner für die Inventarisierung verschiedener Dokumenttypen"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialisiert den DocumentScanner.
        
        Args:
            config: Konfigurationswörterbuch mit Scannereinstellungen
        """
        self.logger = logging.getLogger(__name__)
        self.config = config or {}
        
        # Unterstützte Dateiendungen
        self.supported_extensions = [
            '.pdf', '.docx', '.doc', '.xlsx', '.xls', 
            '.pptx', '.ppt', '.md', '.html', '.txt'
        ]
    
    def scan_directory(self, root_dir: Union[str, Path]) -> pd.DataFrame:
        """
        Scannt ein Verzeichnis nach Dokumenten und extrahiert Metadaten.
        
        Args:
            root_dir: Pfad zum Stammverzeichnis
            
        Returns:
            DataFrame mit Dokumentenmetadaten
        """
        root_dir = Path(root_dir)
        self.logger.info(f"Scanne Verzeichnis: {root_dir}")
        
        document_inventory = []
        
        # Verzeichnis rekursiv durchsuchen
        for file_path in self._find_documents(root_dir):
            try:
                # Basismetadaten extrahieren
                doc_info = self._extract_base_metadata(file_path)
                
                # Dokumentspezifische Metadaten
                if file_path.suffix.lower() == '.pdf':
                    self._extract_pdf_metadata(file_path, doc_info)
                elif file_path.suffix.lower() == '.docx':
                    self._extract_docx_metadata(file_path, doc_info)
                # Weitere Dateitypen hier hinzufügen...
                
                document_inventory.append(doc_info)
                
            except Exception as e:
                self.logger.error(f"Fehler bei {file_path}: {e}", exc_info=True)
        
        if not document_inventory:
            self.logger.warning(f"Keine unterstützten Dokumente in {root_dir} gefunden")
            return pd.DataFrame()
        
        # Konvertiere Liste zu DataFrame
        document_df = pd.DataFrame(document_inventory)
        
        self.logger.info(f"Inventarisierung abgeschlossen: {len(document_df)} Dokumente gefunden")
        return document_df
    
    def _find_documents(self, root_dir: Path) -> List[Path]:
        """Findet alle Dokumente mit unterstützten Dateierweiterungen"""
        documents = []
        
        for file_path in root_dir.glob('**/*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                documents.append(file_path)
        
        return documents
    
    def _extract_base_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extrahiert Basis-Metadaten für alle Dateitypen"""
        file_stats = file_path.stat()
        
        # Erkenne MIME-Typ
        mime_type = magic.from_file(str(file_path), mime=True)
        
        return {
            'filename': file_path.name,
            'path': str(file_path),
            'size_kb': round(file_stats.st_size / 1024, 2),
            'modified_date': datetime.fromtimestamp(file_stats.st_mtime).strftime('%Y-%m-%d'),
            'file_type': mime_type,
            'extension': file_path.suffix.lower(),
            'title': '',
            'author': '',
            'pages': 0,
            'complexity_score': 0,
            'has_tables': False,
            'has_images': False,
            'text_length': 0
        }
    
    def _extract_pdf_metadata(self, file_path: Path, doc_info: Dict[str, Any]) -> None:
        """Extrahiert Metadaten aus PDF-Dokumenten"""
        with fitz.open(str(file_path)) as doc:
            doc_info['pages'] = len(doc)
            metadata = doc.metadata
            if metadata:
                doc_info['title'] = metadata.get('title', '') or file_path.stem
                doc_info['author'] = metadata.get('author', '')
            
            # Komplexitätsanalyse
            has_tables = False
            has_images = False
            text_length = 0
            
            for page in doc:
                # Text extrahieren
                text = page.get_text()
                text_length += len(text)
                
                # Bilder prüfen
                images = page.get_images()
                if images:
                    has_images = True
                
                # Tabellen heuristisch erkennen
                if '|' in text or '\t' in text or (text.count('\n') > 3 and len(text.split('\n')) > 4):
                    # Einfache Heuristik für tabellenähnlichen Text
                    has_tables = True
            
            # Komplexitätsscore berechnen
            complexity = 1  # Basiswert
            if has_tables: complexity += 2
            if has_images: complexity += 1
            if text_length > 10000: complexity += 1
            if doc_info['pages'] > 20: complexity += 1
            
            doc_info['complexity_score'] = complexity
            doc_info['has_tables'] = has_tables
            doc_info['has_images'] = has_images
            doc_info['text_length'] = text_length
    
    def _extract_docx_metadata(self, file_path: Path, doc_info: Dict[str, Any]) -> None:
        """Extrahiert Metadaten aus Word-Dokumenten"""
        doc = docx.Document(str(file_path))
        
        # Metadaten
        core_props = doc.core_properties
        doc_info['title'] = core_props.title or file_path.stem
        doc_info['author'] = core_props.author or ''
        
        # Komplexitätsanalyse
        has_tables = len(doc.tables) > 0
        has_images = sum(1 for p in doc.paragraphs 
                          for r in p.runs 
                          if len(r._element.xpath('.//wp:drawing | .//v:imagedata')) > 0)
        
        paragraphs = len(doc.paragraphs)
        text_length = sum(len(p.text) for p in doc.paragraphs)
        
        # Komplexitätsscore
        complexity = 1  # Basiswert
        if has_tables: complexity += 2
        if has_images: complexity += 1
        if text_length > 10000: complexity += 1
        if paragraphs > 100: complexity += 1
        
        doc_info['complexity_score'] = complexity
        doc_info['has_tables'] = has_tables
        doc_info['has_images'] = has_images
        doc_info['text_length'] = text_length
        doc_info['pages'] = paragraphs // 40 + 1  # Grobe Schätzung


if __name__ == "__main__":
    # Setup Logging
    logging.basicConfig(level=logging.INFO, 
                        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    
    # Test des Scanners
    scanner = DocumentScanner()
    import sys
    
    if len(sys.argv) > 1:
        directory = sys.argv[1]
    else:
        directory = "data/raw_docs"
    
    try:
        df = scanner.scan_directory(directory)
        if not df.empty:
            print(f"Gefundene Dokumente: {len(df)}")
            print("\nDokumenttypen:")
            print(df['file_type'].value_counts())
            
            print("\nKomplexitätsverteilung:")
            print(df['complexity_score'].value_counts().sort_index())
            
            # CSV speichern
            output_file = "doc_converter/data/inventory/document_inventory.csv"
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
            df.to_csv(output_file, index=False)
            print(f"\nInventar gespeichert unter: {output_file}")
    except Exception as e:
        logging.error(f"Fehler beim Scannen: {e}", exc_info=True)
        sys.exit(1)