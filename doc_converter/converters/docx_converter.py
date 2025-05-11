"""
Word-Konverter für die Dokumentenkonvertierungspipeline.
Konvertiert Word-Dokumente (DOCX) in strukturierte Markdown-Dateien.
"""

import os
import re
from pathlib import Path
import logging
import docx
import html
import io
from typing import Dict, Any, List, Tuple, Optional, Set
from PIL import Image
import numpy as np
from base64 import b64encode
import mammoth

from .base_converter import BaseConverter

class DocxConverter(BaseConverter):
    """Konvertiert Word-Dokumente (DOCX) in strukturierte Markdown-Dateien"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialisiert den Word-Konverter.
        
        Args:
            config: Konfigurationswörterbuch mit Konvertereinstellungen
        """
        super().__init__(config)
        
        # Word-spezifische Konfiguration
        self.extract_images = self.config.get('extract_images', True)
        self.max_image_size = self.config.get('max_image_size', 1024)  # Max. Bildbreite oder -höhe in Pixeln
        self.extract_comments = self.config.get('extract_comments', False)
        self.extract_headers_footers = self.config.get('extract_headers_footers', False)
        self.use_mammoth = self.config.get('use_mammoth', True)  # Mammoth für bessere Konvertierung verwenden
    
    def _convert_to_markdown(self, source_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Konvertiert eine Word-Datei in Markdown.
        
        Args:
            source_path: Pfad zur Word-Datei
            
        Returns:
            Tuple aus (markdown_content, metadata)
        """
        self.logger.info(f"Starte Konvertierung von DOCX: {source_path}")
        
        # Verwende zwei Ansätze und wähle den besseren
        markdown_content = ""
        metadata = {}
        
        try:
            # Mammoth-Ansatz (für bessere Struktur)
            if self.use_mammoth:
                self.logger.info("Verwende Mammoth-Konverter für DOCX")
                mammoth_markdown, mammoth_metadata = self._convert_with_mammoth(source_path)
                
                # Aktualisiere Metadaten
                metadata.update(mammoth_metadata)
                
                # Verwende Mammoth-Ausgabe wenn erfolgreich
                if mammoth_markdown:
                    markdown_content = mammoth_markdown
                    self.logger.info("Mammoth-Konvertierung erfolgreich")
                else:
                    self.logger.warning("Mammoth-Konvertierung fehlgeschlagen, verwende python-docx Fallback")
            
            # Wenn Mammoth fehlschlägt oder nicht aktiviert ist, verwende python-docx
            if not markdown_content:
                self.logger.info("Verwende python-docx Konverter")
                docx_markdown, docx_metadata = self._convert_with_python_docx(source_path)
                
                # Aktualisiere Metadaten
                metadata.update(docx_metadata)
                
                # Verwende python-docx Ausgabe
                markdown_content = docx_markdown
                self.logger.info("python-docx Konvertierung erfolgreich")
            
            self.logger.info(f"DOCX-Konvertierung abgeschlossen: {len(markdown_content)} Zeichen")
            
            return markdown_content, metadata
            
        except Exception as e:
            self.logger.error(f"Fehler bei der DOCX-Konvertierung: {e}", exc_info=True)
            # Fallback zu einfachem Text
            try:
                doc = docx.Document(str(source_path))
                text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                
                # Sammle einfache Metadaten
                basic_metadata = {
                    'title': source_path.stem,
                    'original_format': 'DOCX',
                    'error': f"Fehler bei der Konvertierung: {str(e)}"
                }
                
                # Überprüfe core_properties
                try:
                    props = doc.core_properties
                    basic_metadata['title'] = props.title or source_path.stem
                    basic_metadata['author'] = props.author or ''
                    basic_metadata['created'] = str(props.created) if props.created else ''
                    basic_metadata['modified'] = str(props.modified) if props.modified else ''
                except:
                    pass
                
                self.logger.warning("Verwende einfachen Text-Fallback für DOCX")
                return text, basic_metadata
                
            except Exception as fallback_error:
                self.logger.error(f"Auch Fallback-Konvertierung fehlgeschlagen: {fallback_error}", exc_info=True)
                return f"# Konvertierung fehlgeschlagen\n\nFehler: {str(e)}\nFallback-Fehler: {str(fallback_error)}", {
                    'title': source_path.stem,
                    'original_format': 'DOCX',
                    'error': f"Fehler: {str(e)}, Fallback-Fehler: {str(fallback_error)}"
                }
    
    def _convert_with_mammoth(self, source_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Konvertiert eine Word-Datei mit Mammoth.
        
        Args:
            source_path: Pfad zur Word-Datei
            
        Returns:
            Tuple aus (markdown_content, metadata)
        """
        assets_dir = source_path.parent / 'assets'
        os.makedirs(assets_dir, exist_ok=True)
        
        # Bereite Bildkonvertierungsoptionen vor
        image_handler = None
        extracted_images = []
        
        if self.extract_images:
            def handle_image(image, _):
                with image.open() as image_bytes:
                    # Generiere einen eindeutigen Dateinamen
                    image_name = f"image_{len(extracted_images) + 1}"
                    file_ext = '.png'  # Standardwert
                    
                    # Versuche das Bildformat zu bestimmen
                    try:
                        img = Image.open(image_bytes)
                        file_ext = f".{img.format.lower()}" if img.format else '.png'
                    except:
                        pass
                    
                    image_filename = f"{image_name}{file_ext}"
                    image_path = assets_dir / image_filename
                    
                    # Neupositionieren des Dateizeigers
                    image_bytes.seek(0)
                    
                    # Lade und skaliere das Bild
                    try:
                        pil_img = Image.open(image_bytes)
                        width, height = pil_img.size
                        
                        # Skaliere große Bilder herunter
                        if width > self.max_image_size or height > self.max_image_size:
                            scale_factor = min(self.max_image_size / width, self.max_image_size / height)
                            new_width = int(width * scale_factor)
                            new_height = int(height * scale_factor)
                            pil_img = pil_img.resize((new_width, new_height), Image.LANCZOS)
                            
                            pil_img.save(image_path)
                        else:
                            # Speichere das Originalbild
                            with open(image_path, 'wb') as f:
                                image_bytes.seek(0)
                                f.write(image_bytes.read())
                        
                        # Speichere die Bildinformationen
                        extracted_images.append({
                            'path': f"assets/{image_filename}",
                            'alt': image_name
                        })
                        
                        # Generiere Markdown-Bildverweis
                        return {
                            "src": f"assets/{image_filename}",
                            "alt": image_name
                        }
                    except Exception as e:
                        self.logger.warning(f"Fehler beim Verarbeiten des Bildes: {e}")
                        return None
            
            image_handler = handle_image
        
        # Benutzerdefinierte Konvertierungsoptionen
        style_map = """
        p[style-name='Heading 1'] => h1
        p[style-name='Heading 2'] => h2
        p[style-name='Heading 3'] => h3
        p[style-name='Heading 4'] => h4
        p[style-name='Heading 5'] => h5
        p[style-name='heading 1'] => h1
        p[style-name='heading 2'] => h2
        p[style-name='heading 3'] => h3
        p[style-name='heading 4'] => h4
        p[style-name='heading 5'] => h5
        p[style-name='Title'] => h1.title
        p[style-name='Subtitle'] => h2.subtitle
        p[style-name='Quote'] => blockquote
        p[style-name='Intense Quote'] => blockquote.intense
        r[style-name='Strong'] => strong
        r[style-name='Emphasis'] => em
        p[style-name='List Paragraph'] => ul > li:fresh
        table => table
        r[style-name='Hyperlink'] => a
        p[style-name='Normal'] => p:fresh
        """
        
        # Konvertiere das Dokument mit Mammoth
        result = mammoth.convert_to_markdown(
            source_path,
            style_map=style_map,
            convert_image=image_handler
        )
        
        markdown = result.value
        
        # Extrahiere Metadaten
        metadata = self._extract_docx_metadata(source_path)
        
        # Warnungen protokollieren
        for warning in result.messages:
            self.logger.warning(f"Mammoth Warnung: {warning}")
        
        # Führe Nachbearbeitung durch
        markdown = self._postprocess_markdown(markdown)
        
        return markdown, metadata
    
    def _convert_with_python_docx(self, source_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Konvertiert eine Word-Datei mit python-docx.
        
        Args:
            source_path: Pfad zur Word-Datei
            
        Returns:
            Tuple aus (markdown_content, metadata)
        """
        doc = docx.Document(str(source_path))
        metadata = self._extract_docx_metadata(source_path)
        
        assets_dir = source_path.parent / 'assets'
        os.makedirs(assets_dir, exist_ok=True)
        
        # Sammle alle Inhalte
        markdown_parts = []
        
        # Verarbeite Abschnitte und Absätze
        for element in doc.element.body:
            # Verarbeite Tabellen
            if element.tag.endswith('tbl'):
                try:
                    table_markdown = self._convert_table(element)
                    if table_markdown:
                        markdown_parts.append(f"\n{table_markdown}\n")
                except Exception as e:
                    self.logger.warning(f"Fehler bei der Tabellenkonvertierung: {e}")
            
            # Verarbeite Absätze und andere Elemente
            elif element.tag.endswith('p'):
                paragraph = docx.text.paragraph.Paragraph(element, doc)
                markdown_text = self._convert_paragraph(paragraph)
                if markdown_text:
                    markdown_parts.append(markdown_text)
        
        # Extrahiere Bilder, wenn aktiviert
        if self.extract_images:
            try:
                image_data = self._extract_docx_images(doc, assets_dir)
                # Füge Bildverweise an den richtigen Stellen ein
                for image in image_data:
                    # Suche nach einem guten Einfügepunkt für das Bild
                    for i, part in enumerate(markdown_parts):
                        if i > 0 and not part.startswith('#') and len(part.strip()) > 0:
                            # Füge das Bild nach dem ersten nicht-Überschriften-Absatz ein
                            markdown_parts.insert(i + 1, f"\n![{image['alt']}]({image['path']})\n")
                            break
                    else:
                        # Wenn kein guter Einfügepunkt gefunden wurde, füge am Ende ein
                        markdown_parts.append(f"\n![{image['alt']}]({image['path']})\n")
            except Exception as e:
                self.logger.warning(f"Fehler beim Extrahieren von Bildern: {e}")
        
        # Verarbeite Kommentare, wenn aktiviert
        if self.extract_comments and hasattr(doc, 'comments') and doc.comments:
            comments_markdown = self._extract_comments(doc)
            if comments_markdown:
                markdown_parts.append("\n## Kommentare\n")
                markdown_parts.append(comments_markdown)
        
        # Verarbeite Kopf-/Fußzeilen, wenn aktiviert
        if self.extract_headers_footers:
            headers_footers_markdown = self._extract_headers_footers(doc)
            if headers_footers_markdown:
                markdown_parts.append("\n## Kopf- und Fußzeilen\n")
                markdown_parts.append(headers_footers_markdown)
        
        # Kombiniere alle Teile und führe Nachbearbeitung durch
        markdown_content = "\n\n".join(filter(None, markdown_parts))
        
        # Nachbearbeitung
        markdown_content = self._postprocess_markdown(markdown_content)
        
        return markdown_content, metadata
    
    def _convert_paragraph(self, paragraph) -> str:
        """
        Konvertiert einen Word-Absatz in Markdown.
        
        Args:
            paragraph: python-docx Paragraph-Objekt
            
        Returns:
            Markdown-Text
        """
        if not paragraph.text.strip():
            return ""
        
        # Bestimme Absatzstil
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        # Verarbeite Überschriften
        if "heading" in style_name or "überschrift" in style_name:
            level = 1  # Standardmäßig h1
            # Versuche die Ebene aus dem Stil zu extrahieren
            match = re.search(r'\d+', style_name)
            if match:
                level = int(match.group(0))
                level = max(1, min(level, 6))  # Begrenze auf 1-6
            
            return f"{'#' * level} {paragraph.text.strip()}"
        
        # Verarbeite Listen
        elif "list" in style_name or paragraph.paragraph_format.left_indent:
            # Prüfe auf Nummerierung
            if hasattr(paragraph, 'numbering') and paragraph.numbering:
                return f"1. {paragraph.text.strip()}"
            else:
                return f"* {paragraph.text.strip()}"
        
        # Verarbeite Zitate
        elif "quote" in style_name:
            return f"> {paragraph.text.strip()}"
        
        # Verarbeite normalen Text mit Formatierung
        else:
            text = paragraph.text.strip()
            
            # Suche nach Fettdruck und Kursivschrift in Runs
            formatted_text = ""
            for run in paragraph.runs:
                run_text = run.text
                
                # Überpüfe, ob wir Formatierung anwenden müssen
                if run.bold and run.italic:
                    run_text = f"***{run_text}***"
                elif run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                
                formatted_text += run_text
            
            # Wenn wir Formatierung angewendet haben, verwende die formatierte Version
            if formatted_text:
                text = formatted_text
            
            return text
    
    def _convert_table(self, table_element) -> str:
        """
        Konvertiert eine Word-Tabelle in Markdown.
        
        Args:
            table_element: lxml-Element der Tabelle
            
        Returns:
            Markdown-Tabelle
        """
        rows = []
        
        # Finde alle Tabellenzeilen
        row_elements = table_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
        if not row_elements:
            return ""
        
        # Verarbeite jede Zeile
        for row_idx, row_element in enumerate(row_elements):
            row_cells = []
            
            # Finde alle Zellen in der Zeile
            cell_elements = row_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
            
            for cell_element in cell_elements:
                # Extrahiere Text aus der Zelle
                cell_text = ""
                for p in cell_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                    for t in p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        cell_text += t.text if t.text else ""
                
                # Bereinige den Text und füge ihn zur Zelle hinzu
                cell_text = cell_text.strip().replace('\n', ' ')
                row_cells.append(cell_text)
            
            # Baue die Markdown-Tabellenzeile
            if row_cells:
                rows.append('| ' + ' | '.join(row_cells) + ' |')
            
            # Füge nach der ersten Zeile die Trennzeile ein
            if row_idx == 0:
                separator = '| ' + ' | '.join(['---'] * len(row_cells)) + ' |'
                rows.append(separator)
        
        # Leere Tabellen ignorieren
        if len(rows) <= 1:
            return ""
        
        return '\n'.join(rows)
    
    def _extract_docx_images(self, doc, assets_dir: Path) -> List[Dict[str, Any]]:
        """
        Extrahiert Bilder aus einem Word-Dokument.
        
        Args:
            doc: python-docx Document-Objekt
            assets_dir: Verzeichnis für extrahierte Assets
            
        Returns:
            Liste der extrahierten Bilder mit Pfaden und Metadaten
        """
        images = []
        image_count = 0
        
        # Finde die Beziehungen zu den Bildern
        rels = doc.part.rels
        
        for rel_id, rel in rels.items():
            # Nur Bildbeziehungen verarbeiten
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    
                    # Bestimme den Dateityp
                    image_ext = "png"  # Standardwert
                    content_type = image_part.content_type.lower()
                    
                    if "png" in content_type:
                        image_ext = "png"
                    elif "jpeg" in content_type or "jpg" in content_type:
                        image_ext = "jpg"
                    elif "gif" in content_type:
                        image_ext = "gif"
                    elif "tiff" in content_type:
                        image_ext = "tiff"
                    elif "bmp" in content_type:
                        image_ext = "bmp"
                    
                    # Generiere Dateiname für das Bild
                    image_count += 1
                    image_filename = f"image_{image_count}.{image_ext}"
                    image_path = assets_dir / image_filename
                    
                    # Lade und skaliere das Bild
                    try:
                        pil_img = Image.open(io.BytesIO(image_bytes))
                        width, height = pil_img.size
                        
                        # Skaliere große Bilder herunter
                        if width > self.max_image_size or height > self.max_image_size:
                            scale_factor = min(self.max_image_size / width, self.max_image_size / height)
                            new_width = int(width * scale_factor)
                            new_height = int(height * scale_factor)
                            pil_img = pil_img.resize((new_width, new_height), Image.LANCZOS)
                            
                            # Speichere das skalierte Bild
                            pil_img.save(image_path)
                        else:
                            # Speichere das Originalbild
                            with open(image_path, 'wb') as f:
                                f.write(image_bytes)
                        
                        # Füge Bildinformationen zur Liste hinzu
                        images.append({
                            'path': f"assets/{image_filename}",
                            'alt': f"Bild {image_count}",
                            'width': width,
                            'height': height
                        })
                        
                    except Exception as e:
                        self.logger.warning(f"Fehler beim Verarbeiten des Bildes {image_count}: {e}")
                
                except Exception as e:
                    self.logger.warning(f"Fehler beim Extrahieren des Bildes mit rel_id {rel_id}: {e}")
        
        return images
    
    def _extract_comments(self, doc) -> str:
        """
        Extrahiert Kommentare aus einem Word-Dokument.
        
        Args:
            doc: python-docx Document-Objekt
            
        Returns:
            Markdown-Text mit Kommentaren
        """
        try:
            comments_part = doc.part.comments_part
            if not comments_part:
                return ""
            
            comment_elements = comments_part.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment')
            if not comment_elements:
                return ""
            
            comments = []
            for comment in comment_elements:
                author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unbekannt')
                date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                
                # Extrahiere Kommentartext
                comment_text = ""
                for p in comment.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                    for t in p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        comment_text += t.text if t.text else ""
                
                if comment_text.strip():
                    comments.append(f"**{author}** ({date}):\n> {comment_text.strip()}")
            
            return "\n\n".join(comments)
            
        except Exception as e:
            self.logger.warning(f"Fehler beim Extrahieren von Kommentaren: {e}")
            return ""
    
    def _extract_headers_footers(self, doc) -> str:
        """
        Extrahiert Kopf- und Fußzeilen aus einem Word-Dokument.
        
        Args:
            doc: python-docx Document-Objekt
            
        Returns:
            Markdown-Text mit Kopf- und Fußzeilen
        """
        try:
            sections = doc.sections
            header_footer_content = []
            
            for i, section in enumerate(sections):
                section_header = f"### Abschnitt {i+1}"
                header_parts = []
                footer_parts = []
                
                # Verarbeite Kopfzeilen
                if section.header:
                    header_text = '\n'.join(p.text for p in section.header.paragraphs if p.text.strip())
                    if header_text:
                        header_parts.append(f"**Kopfzeile:**\n> {header_text}")
                
                # Verarbeite Fußzeilen
                if section.footer:
                    footer_text = '\n'.join(p.text for p in section.footer.paragraphs if p.text.strip())
                    if footer_text:
                        footer_parts.append(f"**Fußzeile:**\n> {footer_text}")
                
                # Füge zur Gesamtliste hinzu
                if header_parts or footer_parts:
                    header_footer_content.append(section_header)
                    header_footer_content.extend(header_parts)
                    header_footer_content.extend(footer_parts)
            
            return '\n\n'.join(header_footer_content)
            
        except Exception as e:
            self.logger.warning(f"Fehler beim Extrahieren von Kopf-/Fußzeilen: {e}")
            return ""
    
    def _extract_docx_metadata(self, source_path: Path) -> Dict[str, Any]:
        """
        Extrahiert Metadaten aus einem Word-Dokument.
        
        Args:
            source_path: Pfad zur Word-Datei
            
        Returns:
            Dictionary mit Metadaten
        """
        try:
            doc = docx.Document(str(source_path))
            
            # Basis-Metadaten
            metadata = {
                'title': source_path.stem,
                'original_format': 'DOCX',
                'pages': len(doc.paragraphs) // 40 + 1,  # Grobe Schätzung
                'has_tables': bool(doc.tables),
                'has_images': False,  # Wird später aktualisiert
                'has_comments': False  # Wird später aktualisiert
            }
            
            # Erweiterte Metadaten aus Core Properties
            try:
                props = doc.core_properties
                metadata['title'] = props.title or source_path.stem
                metadata['author'] = props.author or ''
                metadata['created'] = str(props.created) if props.created else ''
                metadata['modified'] = str(props.modified) if props.modified else ''
                metadata['subject'] = props.subject or ''
                metadata['keywords'] = props.keywords or ''
                metadata['category'] = props.category or ''
            except Exception as e:
                self.logger.warning(f"Fehler beim Extrahieren der Kern-Metadaten: {e}")
            
            # Prüfe auf Bilder
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        metadata['has_images'] = True
                        break
            except Exception as e:
                self.logger.warning(f"Fehler beim Prüfen auf Bilder: {e}")
            
            # Prüfe auf Kommentare
            try:
                if hasattr(doc.part, 'comments_part') and doc.part.comments_part:
                    metadata['has_comments'] = True
            except Exception as e:
                self.logger.warning(f"Fehler beim Prüfen auf Kommentare: {e}")
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Fehler beim Extrahieren der Metadaten: {e}", exc_info=True)
            return {
                'title': source_path.stem,
                'original_format': 'DOCX',
                'error': f"Fehler bei der Metadatenextraktion: {str(e)}"
            }
    
    def _postprocess_markdown(self, markdown: str) -> str:
        """
        Führt Nachbearbeitung des Markdown-Texts durch.
        
        Args:
            markdown: Ursprünglicher Markdown-Text
            
        Returns:
            Bereinigter und optimierter Markdown-Text
        """
        # Bereinige mehrfache Leerzeilen
        markdown = re.sub(r'\n{3,}', '\n\n', markdown)
        
        # Bereinige Formatierungsprobleme
        markdown = re.sub(r'\*{4,}', '**', markdown)  # Mehr als 3 Sterne in Folge
        
        # Korrigiere falsche Überschriften ohne Leerzeichen
        markdown = re.sub(r'(#{1,6})([^ #])', r'\1 \2', markdown)
        
        # Entferne ungültige HTML-Entitäten
        markdown = re.sub(r'&[^;]+;', '', markdown)
        
        # Entferne Word-spezifische Steuerzeichen
        markdown = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', markdown)
        
        # Decoding von HTML-Entities
        try:
            markdown = html.unescape(markdown)
        except:
            pass
        
        # Korrigiere häufige Probleme mit Tabellen
        # (Leere Zeilen in Tabellen entfernen)
        table_lines = markdown.split('\n')
        in_table = False
        cleaned_lines = []
        
        for line in table_lines:
            if line.strip().startswith('|') and line.strip().endswith('|'):
                in_table = True
                if line.strip():  # Nur nicht-leere Zeilen innerhalb der Tabelle
                    cleaned_lines.append(line)
            elif in_table and not line.strip():
                # Leere Zeile innerhalb einer Tabelle ignorieren
                continue
            else:
                in_table = False
                cleaned_lines.append(line)
        
        markdown = '\n'.join(cleaned_lines)
        
        return markdown


if __name__ == "__main__":
    # Setup Logging
    logging.basicConfig(level=logging.INFO, 
                        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    
    # Test des Word-Konverters
    import sys
    
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        target_dir = sys.argv[2] if len(sys.argv) > 2 else "output"
        
        converter = DocxConverter()
        result = converter.convert(docx_path, target_dir)
        
        if result['success']:
            print(f"Konvertierung erfolgreich: {result['target']}")
        else:
            print(f"Fehler bei der Konvertierung: {result.get('error', 'Unbekannter Fehler')}")
    else:
        print("Verwendung: python docx_converter.py <docx_datei> [ziel_verzeichnis]")