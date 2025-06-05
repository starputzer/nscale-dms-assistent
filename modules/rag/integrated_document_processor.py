"""
Integrierter Document Processor für nahtlose RAG-Pipeline
Verbindet Dokumentenkonvertierung, Chunking und Indexierung
"""
import sys
import os
# Add app directory to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

import asyncio
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import json
from pathlib import Path
import logging
from dataclasses import dataclass, asdict
import aiofiles
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

from ..retrieval.document_store import DocumentStore
from .semantic_chunker import SemanticChunker, Chunk
from .hybrid_retriever import HybridRetriever
from doc_converter.main import DocConverter
from doc_converter.inventory.document_classifier import DocumentClassifier

logger = logging.getLogger(__name__)


@dataclass
class ProcessingResult:
    """Ergebnis der Dokumentenverarbeitung"""
    file_path: str
    status: str  # 'success', 'error', 'skipped'
    chunks_created: int
    processing_time: float
    error_message: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class DocumentMetadata:
    """Erweiterte Metadaten für Dokumente"""
    file_path: str
    file_hash: str
    file_size: int
    document_type: str
    language: str
    created_at: datetime
    modified_at: datetime
    version: int
    tags: List[str]
    extracted_metadata: Dict[str, Any]
    processing_info: Dict[str, Any]


class IntegratedDocumentProcessor:
    """
    Vollständig integrierte Pipeline für Dokumentenverarbeitung
    """
    
    def __init__(self,
                 source_dir: str = "data/raw_docs",
                 converted_dir: str = "data/converted",
                 index_dir: str = "cache/document_index",
                 watch_enabled: bool = True):
        
        self.source_dir = Path(source_dir)
        self.converted_dir = Path(converted_dir)
        self.index_dir = Path(index_dir)
        self.watch_enabled = watch_enabled
        
        # Erstelle Verzeichnisse
        self.source_dir.mkdir(parents=True, exist_ok=True)
        self.converted_dir.mkdir(parents=True, exist_ok=True)
        self.index_dir.mkdir(parents=True, exist_ok=True)
        
        # Komponenten initialisieren
        self.converter = DocConverter()
        self.classifier = DocumentClassifier()
        self.chunker = SemanticChunker(
            min_chunk_size=200,
            target_chunk_size=600,
            max_chunk_size=1000,
            overlap_ratio=0.15
        )
        self.retriever = HybridRetriever(
            index_path=str(self.index_dir)
        )
        
        # Tracking
        self.processed_files = self._load_processed_files()
        self.processing_queue = asyncio.Queue()
        self.observer = None
        
        # Cache für Metadaten
        self.metadata_cache = {}
        
        # Statistiken
        self.stats = {
            'total_processed': 0,
            'total_chunks': 0,
            'total_errors': 0,
            'processing_time': 0.0
        }
    
    async def initialize(self) -> bool:
        """Initialisiert das System und lädt bestehende Indizes"""
        try:
            logger.info("Initialisiere Integrated Document Processor")
            
            # Lade bestehenden Index
            if self.retriever.load_index():
                logger.info("Bestehender Index geladen")
            
            # Starte File Watcher wenn aktiviert
            if self.watch_enabled:
                self._start_file_watcher()
            
            # Starte Background Worker
            asyncio.create_task(self._process_queue_worker())
            
            logger.info("Initialisierung abgeschlossen")
            return True
            
        except Exception as e:
            logger.error(f"Fehler bei der Initialisierung: {e}")
            return False
    
    async def process_document(self, 
                             file_path: str,
                             force_reprocess: bool = False) -> ProcessingResult:
        """
        Verarbeitet ein einzelnes Dokument durch die komplette Pipeline
        """
        start_time = datetime.now()
        file_path = Path(file_path)
        
        try:
            # 1. Prüfe ob bereits verarbeitet
            file_hash = self._calculate_file_hash(file_path)
            if not force_reprocess and file_hash in self.processed_files:
                logger.info(f"Dokument bereits verarbeitet: {file_path}")
                return ProcessingResult(
                    file_path=str(file_path),
                    status='skipped',
                    chunks_created=0,
                    processing_time=0.0,
                    metadata={'reason': 'already_processed'}
                )
            
            # 2. Dokument-Klassifizierung
            doc_type = self.classifier.classify(str(file_path))
            logger.info(f"Dokument klassifiziert als: {doc_type}")
            
            # 3. Metadaten-Extraktion
            metadata = await self._extract_metadata(file_path, doc_type)
            
            # 4. Dokumentenkonvertierung
            converted_path = await self._convert_document(file_path, metadata)
            if not converted_path:
                raise Exception("Konvertierung fehlgeschlagen")
            
            # 5. Text laden
            async with aiofiles.open(converted_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # 6. Intelligentes Chunking
            chunks = self.chunker.chunk_document(
                content,
                metadata={
                    'source': str(file_path),
                    'type': doc_type,
                    'converted_from': file_path.suffix,
                    **metadata.extracted_metadata
                }
            )
            
            logger.info(f"Erstellt {len(chunks)} Chunks für {file_path}")
            
            # 7. Chunks für Indexierung vorbereiten
            index_documents = []
            for i, chunk in enumerate(chunks):
                index_doc = {
                    'text': chunk.text,
                    'metadata': {
                        **chunk.metadata,
                        'chunk_id': f"{file_hash}_{i}",
                        'chunk_index': i,
                        'total_chunks': len(chunks),
                        'coherence_score': chunk.coherence_score,
                        'chunk_type': chunk.chunk_type,
                        'file_path': str(file_path),
                        'file_hash': file_hash,
                        'processed_at': datetime.now().isoformat()
                    }
                }
                index_documents.append(index_doc)
            
            # 8. In Hybrid Retriever indexieren
            await self.retriever.index_documents(index_documents)
            
            # 9. Tracking aktualisieren
            self.processed_files[file_hash] = {
                'file_path': str(file_path),
                'processed_at': datetime.now().isoformat(),
                'chunks': len(chunks),
                'metadata': asdict(metadata)
            }
            self._save_processed_files()
            
            # 10. Statistiken aktualisieren
            processing_time = (datetime.now() - start_time).total_seconds()
            self.stats['total_processed'] += 1
            self.stats['total_chunks'] += len(chunks)
            self.stats['processing_time'] += processing_time
            
            # Cache Metadaten
            self.metadata_cache[file_hash] = metadata
            
            return ProcessingResult(
                file_path=str(file_path),
                status='success',
                chunks_created=len(chunks),
                processing_time=processing_time,
                metadata={
                    'document_type': doc_type,
                    'file_hash': file_hash,
                    'chunk_stats': {
                        'total': len(chunks),
                        'avg_size': sum(c.size for c in chunks) / len(chunks),
                        'avg_coherence': sum(c.coherence_score for c in chunks) / len(chunks)
                    }
                }
            )
            
        except Exception as e:
            logger.error(f"Fehler bei der Verarbeitung von {file_path}: {e}")
            self.stats['total_errors'] += 1
            
            return ProcessingResult(
                file_path=str(file_path),
                status='error',
                chunks_created=0,
                processing_time=(datetime.now() - start_time).total_seconds(),
                error_message=str(e)
            )
    
    async def process_directory(self,
                              directory: Optional[str] = None,
                              pattern: str = "**/*",
                              force_reprocess: bool = False) -> List[ProcessingResult]:
        """
        Verarbeitet alle Dokumente in einem Verzeichnis
        """
        directory = Path(directory) if directory else self.source_dir
        results = []
        
        # Finde alle passenden Dateien
        files = list(directory.glob(pattern))
        logger.info(f"Gefunden: {len(files)} Dateien in {directory}")
        
        # Batch-Verarbeitung
        for file_path in files:
            if file_path.is_file() and file_path.suffix.lower() in [
                '.pdf', '.docx', '.xlsx', '.pptx', '.txt', '.md', '.html'
            ]:
                result = await self.process_document(file_path, force_reprocess)
                results.append(result)
        
        return results
    
    async def _extract_metadata(self, 
                              file_path: Path, 
                              doc_type: str) -> DocumentMetadata:
        """Extrahiert erweiterte Metadaten"""
        file_stat = file_path.stat()
        
        # Basis-Metadaten
        metadata = DocumentMetadata(
            file_path=str(file_path),
            file_hash=self._calculate_file_hash(file_path),
            file_size=file_stat.st_size,
            document_type=doc_type,
            language='de',  # TODO: Automatische Spracherkennung
            created_at=datetime.fromtimestamp(file_stat.st_ctime),
            modified_at=datetime.fromtimestamp(file_stat.st_mtime),
            version=1,
            tags=[],
            extracted_metadata={},
            processing_info={}
        )
        
        # Format-spezifische Metadaten
        if file_path.suffix.lower() == '.pdf':
            # PDF-Metadaten mit PyMuPDF extrahieren
            try:
                import fitz  # PyMuPDF
                with fitz.open(str(file_path)) as pdf:
                    pdf_metadata = pdf.metadata
                    metadata.extracted_metadata.update({
                        'title': pdf_metadata.get('title', ''),
                        'author': pdf_metadata.get('author', ''),
                        'subject': pdf_metadata.get('subject', ''),
                        'keywords': pdf_metadata.get('keywords', ''),
                        'pages': pdf.page_count,
                        'producer': pdf_metadata.get('producer', ''),
                        'creator': pdf_metadata.get('creator', '')
                    })
            except Exception as e:
                logger.warning(f"PDF-Metadaten konnten nicht extrahiert werden: {e}")
        
        elif file_path.suffix.lower() == '.docx':
            # DOCX-Metadaten
            try:
                from docx import Document
                doc = Document(str(file_path))
                core_props = doc.core_properties
                metadata.extracted_metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'created': core_props.created.isoformat() if core_props.created else '',
                    'modified': core_props.modified.isoformat() if core_props.modified else '',
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables)
                })
            except Exception as e:
                logger.warning(f"DOCX-Metadaten konnten nicht extrahiert werden: {e}")
        
        # Auto-Tagging basierend auf Inhalt/Pfad
        metadata.tags = self._generate_auto_tags(file_path, doc_type)
        
        return metadata
    
    async def _convert_document(self, 
                              file_path: Path, 
                              metadata: DocumentMetadata) -> Optional[Path]:
        """Konvertiert Dokument zu Markdown"""
        try:
            # Ziel-Pfad
            relative_path = file_path.relative_to(self.source_dir)
            target_path = self.converted_dir / relative_path.with_suffix('.md')
            target_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Konvertierung durchführen
            # TODO: Integration mit doc_converter module
            # Für jetzt: Kopiere einfach Text-Dateien
            if file_path.suffix.lower() in ['.txt', '.md']:
                import shutil
                shutil.copy2(file_path, target_path)
            else:
                # Placeholder für echte Konvertierung
                async with aiofiles.open(target_path, 'w', encoding='utf-8') as f:
                    await f.write(f"# {file_path.name}\n\n")
                    await f.write(f"Konvertiert aus: {file_path.suffix}\n\n")
                    await f.write("(Konvertierungs-Placeholder)\n")
            
            return target_path
            
        except Exception as e:
            logger.error(f"Fehler bei der Konvertierung: {e}")
            return None
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Berechnet SHA256-Hash der Datei"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _generate_auto_tags(self, file_path: Path, doc_type: str) -> List[str]:
        """Generiert automatische Tags basierend auf Datei-Eigenschaften"""
        tags = [doc_type]
        
        # Pfad-basierte Tags
        parts = file_path.parts
        if 'handbuch' in str(file_path).lower():
            tags.append('manual')
        if 'anleitung' in str(file_path).lower():
            tags.append('guide')
        if 'admin' in str(file_path).lower():
            tags.append('administration')
        
        # Größen-basierte Tags
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 10:
            tags.append('large')
        elif size_mb < 0.1:
            tags.append('small')
        
        return tags
    
    def _load_processed_files(self) -> Dict[str, Any]:
        """Lädt Liste der bereits verarbeiteten Dateien"""
        tracking_file = self.index_dir / "processed_files.json"
        if tracking_file.exists():
            with open(tracking_file, 'r') as f:
                return json.load(f)
        return {}
    
    def _save_processed_files(self):
        """Speichert Liste der verarbeiteten Dateien"""
        tracking_file = self.index_dir / "processed_files.json"
        with open(tracking_file, 'w') as f:
            json.dump(self.processed_files, f, indent=2)
    
    def _start_file_watcher(self):
        """Startet File System Watcher für automatische Verarbeitung"""
        class DocumentHandler(FileSystemEventHandler):
            def __init__(self, processor):
                self.processor = processor
                
            def on_created(self, event):
                if not event.is_directory:
                    asyncio.create_task(
                        self.processor.processing_queue.put(event.src_path)
                    )
            
            def on_modified(self, event):
                if not event.is_directory:
                    asyncio.create_task(
                        self.processor.processing_queue.put(event.src_path)
                    )
        
        self.observer = Observer()
        self.observer.schedule(
            DocumentHandler(self),
            str(self.source_dir),
            recursive=True
        )
        self.observer.start()
        logger.info(f"File Watcher gestartet für: {self.source_dir}")
    
    async def _process_queue_worker(self):
        """Background Worker für Queue-Verarbeitung"""
        while True:
            try:
                file_path = await self.processing_queue.get()
                logger.info(f"Verarbeite aus Queue: {file_path}")
                await self.process_document(file_path)
            except Exception as e:
                logger.error(f"Fehler im Queue Worker: {e}")
            await asyncio.sleep(0.1)
    
    async def search(self,
                    query: str,
                    k: int = 10,
                    filters: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """
        Durchsucht die indizierten Dokumente
        """
        results = await self.retriever.search(
            query=query,
            k=k,
            use_reranking=True,
            filters=filters
        )
        
        # Erweitere Ergebnisse mit Metadaten
        enhanced_results = []
        for result in results:
            enhanced = result.to_dict()
            
            # Füge Datei-Metadaten hinzu wenn verfügbar
            file_hash = result.metadata.get('file_hash')
            if file_hash and file_hash in self.metadata_cache:
                enhanced['file_metadata'] = asdict(self.metadata_cache[file_hash])
            
            enhanced_results.append(enhanced)
        
        return enhanced_results
    
    def get_statistics(self) -> Dict[str, Any]:
        """Gibt Statistiken über die Verarbeitung zurück"""
        retriever_stats = self.retriever.get_statistics()
        
        return {
            'processor': {
                'total_processed': self.stats['total_processed'],
                'total_chunks': self.stats['total_chunks'],
                'total_errors': self.stats['total_errors'],
                'avg_processing_time': (
                    self.stats['processing_time'] / self.stats['total_processed']
                    if self.stats['total_processed'] > 0 else 0
                ),
                'tracked_files': len(self.processed_files)
            },
            'retriever': retriever_stats,
            'directories': {
                'source': str(self.source_dir),
                'converted': str(self.converted_dir),
                'index': str(self.index_dir)
            }
        }
    
    async def reindex_all(self, force: bool = False) -> List[ProcessingResult]:
        """
        Reindexiert alle Dokumente
        """
        logger.info("Starte vollständige Reindexierung")
        
        if force:
            # Lösche bestehenden Index
            self.processed_files = {}
            self._save_processed_files()
            
        # Verarbeite alle Dokumente
        results = await self.process_directory(force_reprocess=force)
        
        logger.info(f"Reindexierung abgeschlossen: {len(results)} Dokumente")
        return results


# Beispiel-Verwendung
async def example_usage():
    """Beispiel für die Verwendung des Integrated Document Processor"""
    
    # Processor initialisieren
    processor = IntegratedDocumentProcessor(
        source_dir="data/raw_docs",
        converted_dir="data/converted",
        watch_enabled=True
    )
    
    # System initialisieren
    await processor.initialize()
    
    # Einzelnes Dokument verarbeiten
    result = await processor.process_document("data/raw_docs/handbuch.pdf")
    print(f"Verarbeitung: {result.status}, Chunks: {result.chunks_created}")
    
    # Verzeichnis verarbeiten
    results = await processor.process_directory()
    print(f"Verarbeitet: {len(results)} Dokumente")
    
    # Suche durchführen
    search_results = await processor.search("Installation von nscale", k=5)
    for i, result in enumerate(search_results):
        print(f"\nErgebnis {i+1}:")
        print(f"  Score: {result['score']:.3f}")
        print(f"  Text: {result['text'][:100]}...")
        print(f"  Quelle: {result['metadata'].get('source', 'unknown')}")
    
    # Statistiken anzeigen
    stats = processor.get_statistics()
    print(f"\nStatistiken: {json.dumps(stats, indent=2)}")


if __name__ == "__main__":
    asyncio.run(example_usage())